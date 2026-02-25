"""Convert ENGR-KB Word documents to KB package format.

Reads each .docx, splits on H1 headings, and generates standard
KB packages: _manifest.json, theory.md, equations.json, tables.json, standards.md.
"""

from __future__ import annotations

import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

# ---------------------------------------------------------------------------
# Configuration: one entry per doc → discipline + package mapping
# ---------------------------------------------------------------------------

DOC_CONFIG = {
    "ENGR-KB_02_Cost_Estimating - Copy.docx": {
        "discipline": "construction-estimating",
        "skip_sections": {"Table of Contents"},
        "overview_section": "1. Cost Estimating Domain Overview",
        "eq_catalog_sections": {
            "9. Complete Equation Catalog",
            "10. Reference Data Tables Summary",
        },
    },
    "ENGR-KB_03_Specifications.docx": {
        "discipline": "specifications",
        "skip_sections": {"Table of Contents"},
        "overview_section": "1. Specifications Domain Overview",
        "eq_catalog_sections": set(),
    },
    "ENGR-KB_04_Project_Management.docx": {
        "discipline": "project-management",
        "skip_sections": {"Table of Contents"},
        "overview_section": "1. Project Management Domain Overview",
        "eq_catalog_sections": {"8. Equation Catalog & Reference Data"},
    },
    "ENGR-KB_05_Construction_Engineering.docx": {
        "discipline": "construction-engineering",
        "skip_sections": {"Table of Contents"},
        "overview_section": "1. Construction Engineering Domain Overview",
        "eq_catalog_sections": {"7. Equation Catalog & Reference Data"},
    },
    "ENGR-KB_06_Cross_Domain_Integration.docx": {
        "discipline": "cross-domain-integration",
        "skip_sections": {"Table of Contents"},
        "overview_section": "1. Cross-Domain Integration Overview",
        "eq_catalog_sections": {"6. System Summary & Master Equation Catalog"},
    },
}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def slugify(text: str) -> str:
    """Convert heading text to URL-safe slug."""
    # Remove leading number prefix like "2. " or "2.1 "
    text = re.sub(r"^\d+(\.\d+)*\.?\s*", "", text)
    slug = text.lower().strip()
    slug = re.sub(r"[^a-z0-9]+", "-", slug)
    slug = slug.strip("-")
    return slug[:60]


def para_to_md(para) -> str:
    """Convert a paragraph to markdown."""
    style = para.style.name if para.style else ""
    text = para.text

    if style == "Heading 2":
        return f"## {text}\n"
    elif style == "Heading 3":
        return f"### {text}\n"
    elif style == "Heading 4":
        return f"#### {text}\n"
    elif style == "List Paragraph":
        return f"- {text}\n"
    elif text.strip():
        return f"{text}\n"
    return ""


def table_to_md(tbl) -> str:
    """Convert a Word table to markdown table."""
    rows = []
    for row in tbl.rows:
        cells = [c.text.strip().replace("\n", " ") for c in row.cells]
        rows.append(cells)
    if not rows:
        return ""

    lines = []
    # Header
    lines.append("| " + " | ".join(rows[0]) + " |")
    lines.append("| " + " | ".join(["---"] * len(rows[0])) + " |")
    for row in rows[1:]:
        # Pad if needed
        while len(row) < len(rows[0]):
            row.append("")
        lines.append("| " + " | ".join(row[: len(rows[0])]) + " |")
    return "\n".join(lines) + "\n"


def table_to_json(tbl, table_id: str, title: str = "") -> dict:
    """Convert a Word table to the KB tables.json format."""
    rows = []
    for row in tbl.rows:
        cells = [c.text.strip().replace("\n", " ") for c in row.cells]
        rows.append(cells)
    if not rows:
        return None

    headers = rows[0]
    data_rows = rows[1:]

    return {
        "id": table_id,
        "title": title or "Reference Table",
        "type": "reference",
        "data": {"headers": headers, "rows": data_rows},
        "source": "ENGR-KB",
    }


def extract_equations_from_table(tbl, id_prefix: str, pkg_slug: str = "") -> list[dict]:
    """Extract equations from a catalog-style table (ID|Name|Formula|...).

    If pkg_slug is provided, equation IDs are scoped to the package
    to avoid cross-package collisions.
    """
    rows = []
    for row in tbl.rows:
        cells = [c.text.strip() for c in row.cells]
        rows.append(cells)
    if len(rows) < 2:
        return []

    headers = [h.lower() for h in rows[0]]

    # Find column indices
    id_col = None
    name_col = None
    formula_col = None
    source_col = None

    for i, h in enumerate(headers):
        if h in ("id", "eq id"):
            id_col = i
        elif h in ("name", "metric"):
            name_col = i
        elif "formula" in h:
            formula_col = i
        elif h in ("source", "type", "category"):
            source_col = i

    if id_col is None or formula_col is None:
        return []

    equations = []
    for row in rows[1:]:
        if len(row) <= max(id_col, formula_col):
            continue
        eq_id = row[id_col].strip()
        if not eq_id:
            continue

        # Normalize the ID — scope to package if provided
        raw_id = eq_id.lower().replace(" ", "-")
        if pkg_slug:
            eq_id_normalized = f"{id_prefix}.{pkg_slug}.{raw_id}"
        else:
            eq_id_normalized = f"{id_prefix}.{raw_id}"

        eq = {
            "id": eq_id_normalized,
            "title": row[name_col].strip() if name_col is not None and len(row) > name_col else eq_id,
            "equation": row[formula_col].strip(),
            "variables": {},
            "applicability": "",
            "source": row[source_col].strip() if source_col is not None and len(row) > source_col else "ENGR-KB",
        }
        equations.append(eq)

    return equations


def is_equation_catalog_table(tbl) -> bool:
    """Check if a table looks like an equation catalog table."""
    if len(tbl.rows) < 2:
        return False
    headers = [c.text.strip().lower() for c in tbl.rows[0].cells]
    has_id = any(h in ("id", "eq id") for h in headers)
    has_formula = any("formula" in h for h in headers)
    # Check first data row has an ID-like value
    if has_id and has_formula and len(tbl.rows) > 1:
        first_id = tbl.rows[1].cells[0].text.strip()
        if re.match(r"[A-Z]{2,}", first_id):
            return True
    return False


def extract_standards_from_text(text: str) -> list[str]:
    """Extract standard references (ASTM, AASHTO, AACE, ACI, etc.) from text."""
    patterns = [
        r"ASTM\s+[A-Z]\d+[\w-]*",
        r"AASHTO\s+[A-Z][\w.-]+",
        r"AACE\s+(?:RP\s+)?[\d]+[RS]-\d+",
        r"ACI\s+\d+[\w.-]*",
        r"USACE\s+E[PM]\s+[\d-]+",
        r"OSHA\s+\d+\s+CFR\s+[\d.]+",
        r"CSI\s+MasterFormat",
    ]
    refs = set()
    for pattern in patterns:
        for m in re.finditer(pattern, text):
            refs.add(m.group())
    return sorted(refs)


# ---------------------------------------------------------------------------
# Document body iteration (preserving paragraph/table order)
# ---------------------------------------------------------------------------


def iter_body_elements(doc):
    """Yield ('para', paragraph) or ('table', table) in document order."""
    body = doc.element.body
    para_idx = 0
    tbl_idx = 0

    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            if para_idx < len(doc.paragraphs):
                yield ("para", doc.paragraphs[para_idx])
                para_idx += 1
        elif tag == "tbl":
            if tbl_idx < len(doc.tables):
                yield ("table", doc.tables[tbl_idx])
                tbl_idx += 1


# ---------------------------------------------------------------------------
# Main conversion
# ---------------------------------------------------------------------------


def split_into_sections(doc) -> list[dict]:
    """Split document into H1 sections with their content."""
    sections = []
    current = None

    for kind, element in iter_body_elements(doc):
        if kind == "para":
            style = element.style.name if element.style else ""
            if style == "Heading 1":
                if current:
                    sections.append(current)
                current = {
                    "title": element.text.strip(),
                    "paragraphs": [],
                    "tables": [],
                }
                continue

        if current is None:
            continue

        if kind == "para":
            current["paragraphs"].append(element)
        elif kind == "table":
            current["tables"].append(element)

    if current:
        sections.append(current)

    return sections


def section_to_theory_md(section: dict) -> str:
    """Convert a section's paragraphs and tables to theory markdown."""
    parts = []

    # We need to interleave paragraphs and tables in order
    # For simplicity, convert paragraphs first, then insert table markers
    # Actually, let's re-iterate the body to maintain order

    # Simple approach: all paragraphs as markdown, tables inserted inline
    # Since we stored them separately, output paragraphs then tables
    for para in section["paragraphs"]:
        md = para_to_md(para)
        if md:
            parts.append(md)

    # Add tables inline as markdown
    for tbl in section["tables"]:
        if not is_equation_catalog_table(tbl):
            parts.append("\n" + table_to_md(tbl) + "\n")

    text = "\n".join(parts)
    # Clean up excessive newlines
    text = re.sub(r"\n{4,}", "\n\n\n", text)
    return text.strip() + "\n"


def process_doc(fname: str, config: dict, output_root: Path):
    """Process one Word document into KB packages."""
    print(f"\n{'='*60}")
    print(f"Processing: {fname}")
    print(f"{'='*60}")

    doc = Document(fname)
    sections = split_into_sections(doc)
    discipline = config["discipline"]
    skip = config["skip_sections"]
    overview_title = config.get("overview_section", "")
    eq_catalog_titles = config.get("eq_catalog_sections", set())

    # Collect all equation catalog equations
    all_catalog_eqs = []
    for sec in sections:
        if sec["title"] in eq_catalog_titles:
            for tbl in sec["tables"]:
                if is_equation_catalog_table(tbl):
                    eqs = extract_equations_from_table(tbl, discipline)
                    all_catalog_eqs.extend(eqs)
            # Also check for variable description tables in the section
            # that aren't equation catalogs — these become reference tables

    print(f"  Found {len(all_catalog_eqs)} equations in catalog tables")

    # Determine which sections become packages
    package_sections = []
    overview_section = None

    for sec in sections:
        if sec["title"] in skip or sec["title"] in eq_catalog_titles:
            continue
        if sec["title"] == overview_title:
            overview_section = sec
            continue
        package_sections.append(sec)

    # Assign catalog equations to packages by best-match prefix
    eq_assignment = assign_equations_to_packages(
        all_catalog_eqs, package_sections, discipline
    )

    # Build each package
    for sec in package_sections:
        slug = slugify(sec["title"])
        pkg_id = f"{discipline}.{slug}"
        pkg_dir = output_root / slug
        pkg_dir.mkdir(parents=True, exist_ok=True)

        # Theory: section content + overview preamble for the first package
        theory_parts = []
        if overview_section and sec == package_sections[0]:
            theory_parts.append(f"# {overview_section['title']}\n\n")
            theory_parts.append(section_to_theory_md(overview_section))
            theory_parts.append(f"\n---\n\n# {sec['title']}\n\n")
        else:
            theory_parts.append(f"# {sec['title']}\n\n")
        theory_parts.append(section_to_theory_md(sec))
        theory_md = "".join(theory_parts)

        # Equations: from catalog + any inline equation tables
        # Re-scope catalog equation IDs to this package to avoid collisions
        raw_eqs = eq_assignment.get(sec["title"], [])
        pkg_eqs = []
        for eq in raw_eqs:
            scoped = dict(eq)
            # Replace discipline-level ID with package-scoped ID
            old_id = scoped["id"]
            # Extract the raw suffix (e.g., "pm-020") from "project-management.pm-020"
            raw_suffix = old_id.split(".", 1)[-1] if "." in old_id else old_id
            scoped["id"] = f"{pkg_id}.{raw_suffix}"
            pkg_eqs.append(scoped)

        # Also check inline tables in this section
        for tbl in sec["tables"]:
            if is_equation_catalog_table(tbl):
                inline_eqs = extract_equations_from_table(tbl, discipline, slug)
                # Avoid duplicates
                existing_ids = {e["id"] for e in pkg_eqs}
                for eq in inline_eqs:
                    if eq["id"] not in existing_ids:
                        pkg_eqs.append(eq)

        # Tables: non-equation tables from this section
        pkg_tables = []
        tbl_counter = 1
        for tbl in sec["tables"]:
            if is_equation_catalog_table(tbl):
                continue
            tbl_id = f"{pkg_id}.tbl-{tbl_counter:03d}"
            # Try to get title from context
            tbl_title = ""
            headers = [c.text.strip() for c in tbl.rows[0].cells]
            tbl_title = " / ".join(headers[:3])

            tbl_json = table_to_json(tbl, tbl_id, tbl_title)
            if tbl_json:
                pkg_tables.append(tbl_json)
                tbl_counter += 1

        # Standards: extract from theory text
        all_text = theory_md
        for tbl in sec["tables"]:
            for row in tbl.rows:
                for cell in row.cells:
                    all_text += " " + cell.text
        std_refs = extract_standards_from_text(all_text)

        standards_md = f"## Standards References — {sec['title']}\n\n"
        if std_refs:
            for ref in std_refs:
                standards_md += f"- **{ref}**\n"
        else:
            standards_md += "No specific standards referenced in this section.\n"

        # Manifest
        manifest = {
            "id": pkg_id,
            "name": re.sub(r"^\d+(\.\d+)*\.?\s*", "", sec["title"]).strip(),
            "level": 2,
            "discipline": discipline,
            "parent": discipline,
            "dependencies": [],
            "tags": extract_tags(sec),
            "standards": std_refs[:10],
            "content_files": [
                "theory.md",
                "equations.json",
                "tables.json",
                "standards.md",
            ],
            "equation_count": len(pkg_eqs),
            "table_count": len(pkg_tables),
            "concept_count": count_h2(sec),
            "enabled": True,
            "version": "1.0",
            "last_updated": "2026-02-25",
        }

        # Write files
        (pkg_dir / "theory.md").write_text(theory_md, encoding="utf-8")
        (pkg_dir / "equations.json").write_text(
            json.dumps(pkg_eqs, indent=2, ensure_ascii=False) + "\n", encoding="utf-8"
        )
        (pkg_dir / "tables.json").write_text(
            json.dumps(pkg_tables, indent=2, ensure_ascii=False) + "\n", encoding="utf-8"
        )
        (pkg_dir / "standards.md").write_text(standards_md, encoding="utf-8")
        (pkg_dir / "_manifest.json").write_text(
            json.dumps(manifest, indent=2, ensure_ascii=False) + "\n", encoding="utf-8"
        )

        print(
            f"  {slug}: {len(pkg_eqs)} eqs, {len(pkg_tables)} tbls, "
            f"{len(std_refs)} stds, {len(theory_md)} chars"
        )


def assign_equations_to_packages(
    equations: list[dict], sections: list[dict], discipline: str
) -> dict[str, list[dict]]:
    """Assign catalog equations to package sections by ID prefix matching."""
    assignment: dict[str, list[dict]] = {s["title"]: [] for s in sections}

    # Build keyword mapping from section titles
    section_keywords = {}
    for sec in sections:
        title_lower = sec["title"].lower()
        keywords = set()
        # Extract meaningful words
        for word in re.findall(r"[a-z]+", title_lower):
            if len(word) > 3:
                keywords.add(word)
        section_keywords[sec["title"]] = keywords

    # Known ID prefix → section keyword mappings
    prefix_hints = {
        "esc": ["escalation", "location"],
        "loc": ["escalation", "location"],
        "par": ["parametric"],
        "ctg": ["contingency"],
        "ew": ["earthwork"],
        "lab": ["labor", "markup", "rate"],
        "mkp": ["labor", "markup", "rate"],
        "eqp": ["equipment", "labor", "rate"],
        "cpm": ["cpm", "scheduling", "schedule"],
        "evm": ["earned", "value"],
        "rsk": ["risk"],
        "chg": ["change"],
        "res": ["resource"],
        "del": ["delay"],
        "cf": ["cash", "cost", "loaded", "schedule"],
        "rpt": ["reporting", "output"],
        "agt": ["agent", "interface"],
    }

    for eq in equations:
        eq_id = eq["id"].split(".")[-1]  # e.g., "ce-esc-001"
        parts = eq_id.split("-")
        assigned = False

        # Try prefix hint matching
        for part in parts[1:-1]:  # skip domain prefix and number
            if part in prefix_hints:
                hint_keywords = prefix_hints[part]
                for sec_title, sec_kws in section_keywords.items():
                    if any(hk in sec_kws or hk in sec_title.lower() for hk in hint_keywords):
                        assignment[sec_title].append(eq)
                        assigned = True
                        break
            if assigned:
                break

        # Fallback: assign to first section if unmatched
        if not assigned and sections:
            # Try matching equation title words to section keywords
            eq_words = set(re.findall(r"[a-z]+", eq.get("title", "").lower()))
            best_match = None
            best_score = 0
            for sec_title, sec_kws in section_keywords.items():
                score = len(eq_words & sec_kws)
                if score > best_score:
                    best_score = score
                    best_match = sec_title
            if best_match and best_score > 0:
                assignment[best_match].append(eq)
            elif sections:
                assignment[sections[0]["title"]].append(eq)

    return assignment


def extract_tags(section: dict) -> list[str]:
    """Extract meaningful tags from a section."""
    all_text = " ".join(p.text for p in section["paragraphs"] if p.text)
    title_text = section["title"]
    combined = f"{title_text} {all_text}"

    # Look for domain-specific terms
    tag_patterns = [
        r"\bAACE\b", r"\bCPM\b", r"\bEVM\b", r"\bMasterFormat\b",
        r"\bRSMeans\b", r"\bmonte\s+carlo\b", r"\bPERT\b",
        r"\bearned\s+value\b", r"\bcritical\s+path\b",
        r"\bcontingency\b", r"\bescalation\b", r"\bparametric\b",
        r"\bunit\s+cost\b", r"\bcrew\b", r"\bassembly\b",
        r"\bearthwork\b", r"\bexcavat\w+\b", r"\bcompact\w+\b",
        r"\bdewater\w+\b", r"\bpile\s+driv\w+\b", r"\bdrilled\s+shaft\b",
        r"\bconcrete\b", r"\bgrout\w*\b", r"\bshoring\b",
        r"\bchange\s+order\b", r"\brisk\b", r"\bschedul\w+\b",
    ]
    tags = set()
    for pattern in tag_patterns:
        if re.search(pattern, combined, re.IGNORECASE):
            tag = re.sub(r"\\[bws+*]", "", pattern).replace("\\", "").strip()
            tags.add(tag.lower().replace(" ", "-"))

    # Add from heading
    for word in re.findall(r"[A-Za-z]+", title_text):
        if len(word) > 4:
            tags.add(word.lower())

    return sorted(tags)[:15]


def count_h2(section: dict) -> int:
    """Count H2 headings in a section (concept count)."""
    return sum(
        1
        for p in section["paragraphs"]
        if p.style and p.style.name == "Heading 2"
    )


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main():
    src_dir = Path(".")
    output_root = Path("packages")
    output_root.mkdir(exist_ok=True)

    for fname, config in DOC_CONFIG.items():
        fpath = src_dir / fname
        if not fpath.exists():
            print(f"  SKIP (not found): {fname}")
            continue

        discipline = config["discipline"]
        disc_dir = output_root / discipline
        disc_dir.mkdir(parents=True, exist_ok=True)

        process_doc(str(fpath), config, disc_dir)

    # Summary
    print(f"\n{'='*60}")
    print("SUMMARY")
    print(f"{'='*60}")
    total_pkgs = 0
    total_eqs = 0
    total_tbls = 0
    for disc_dir in sorted(output_root.iterdir()):
        if not disc_dir.is_dir():
            continue
        for pkg_dir in sorted(disc_dir.iterdir()):
            if not pkg_dir.is_dir():
                continue
            mf = pkg_dir / "_manifest.json"
            if mf.exists():
                data = json.loads(mf.read_text(encoding="utf-8"))
                total_pkgs += 1
                total_eqs += data.get("equation_count", 0)
                total_tbls += data.get("table_count", 0)
                print(
                    f"  {data['id']:55s} {data['equation_count']:3d} eqs  "
                    f"{data['table_count']:3d} tbls"
                )
    print(f"\nTotal: {total_pkgs} packages, {total_eqs} equations, {total_tbls} tables")


if __name__ == "__main__":
    main()
